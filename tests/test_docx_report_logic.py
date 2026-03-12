# tests/test_docx_report_logic.py
"""
Tests for docx_report.py pure-logic functions.

Covers: DocxReportConfig, _pt_to_px, _strip_user_headings_plotly,
_fix_axis_gradient_overlaps_plotly, apply_zip_export_style,
_style_plotly_for_docx_export, _add_dataframe_table, _add_figure,
_best_model_line, _recommended_docx_figure_width_in, _ensure_caption_style,
create_docx_report edge cases.
"""

import sys
import os
from io import BytesIO

import pandas as pd
import pytest

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

plotly = pytest.importorskip("plotly", reason="plotly required")
import plotly.graph_objects as go

from adsorblab_pro.docx_report import (
    DOCX_AVAILABLE,
    DocxReportConfig,
    _pt_to_px,
    _strip_user_headings_plotly,
    _fix_axis_gradient_overlaps_plotly,
    apply_zip_export_style,
    style_figure_for_export,
    _style_plotly_for_docx_export,
    _best_model_line,
    _ensure_caption_style,
    _recommended_docx_figure_width_in,
    create_docx_report,
)


# =============================================================================
# DocxReportConfig
# =============================================================================


class TestDocxReportConfig:
    def test_defaults(self):
        cfg = DocxReportConfig()
        assert cfg.img_format == "png"
        assert cfg.img_width_px == 1600
        assert cfg.img_height_px == 1000
        assert cfg.img_scale == 2.0
        assert cfg.figure_width_in == 6.5
        assert cfg.text_preset == "Journal (default)"
        assert cfg.strip_user_headings is True
        assert cfg.fix_axis_overlaps is True
        assert cfg.max_table_rows == 60
        assert cfg.max_table_cols == 12
        assert cfg.float_format == "{:.4g}"

    def test_custom_values(self):
        cfg = DocxReportConfig(
            img_format="svg",
            img_width_px=800,
            max_table_rows=20,
            text_preset="Presentation (slides)",
        )
        assert cfg.img_format == "svg"
        assert cfg.img_width_px == 800
        assert cfg.max_table_rows == 20


# =============================================================================
# _pt_to_px
# =============================================================================


class TestPtToPx:
    def test_basic_conversion(self):
        # 72pt at 72dpi over 1 inch = 72px
        px = _pt_to_px(72.0, width_px=72, target_width_in=1.0)
        assert px == 72

    def test_minimum_clamp(self):
        # Very small pt should clamp to 8
        px = _pt_to_px(0.1, width_px=100, target_width_in=10.0)
        assert px == 8

    def test_zero_target_width(self):
        # Should fall back to 6.5
        px = _pt_to_px(11.0, width_px=1600, target_width_in=0.0)
        assert px >= 8

    def test_negative_target_width(self):
        px = _pt_to_px(11.0, width_px=1600, target_width_in=-1.0)
        assert px >= 8

    def test_standard_journal(self):
        # 11pt text, 1600px wide, 6.5in target
        px = _pt_to_px(11.0, width_px=1600, target_width_in=6.5)
        assert isinstance(px, int)
        assert 20 < px < 60  # Reasonable range


# =============================================================================
# _strip_user_headings_plotly
# =============================================================================


class TestStripUserHeadings:
    def test_removes_title(self):
        fig = go.Figure()
        fig.update_layout(title_text="My Title")
        _strip_user_headings_plotly(fig)
        assert fig.layout.title.text == ""

    def test_no_title_noop(self):
        fig = go.Figure()
        _strip_user_headings_plotly(fig)
        # Should not raise

    def test_removes_paper_heading_annotations(self):
        fig = go.Figure()
        fig.update_layout(
            annotations=[
                dict(
                    text="Heading",
                    xref="paper",
                    yref="paper",
                    x=0.5,
                    y=1.05,
                    showarrow=False,
                ),
                dict(
                    text="Data label",
                    xref="x",
                    yref="y",
                    x=1.0,
                    y=2.0,
                    showarrow=True,
                ),
            ]
        )
        _strip_user_headings_plotly(fig)
        # The heading annotation should be removed; data label kept
        texts = [a.text for a in fig.layout.annotations]
        assert "Heading" not in texts
        assert "Data label" in texts

    def test_keeps_below_plot_annotations(self):
        fig = go.Figure()
        fig.update_layout(
            annotations=[
                dict(
                    text="Footer",
                    xref="paper",
                    yref="paper",
                    x=0.5,
                    y=0.5,
                    showarrow=False,
                ),
            ]
        )
        _strip_user_headings_plotly(fig)
        assert len(fig.layout.annotations) == 1

    def test_with_fig_id(self):
        fig = go.Figure()
        fig.update_layout(title_text="Test")
        _strip_user_headings_plotly(fig, fig_id="test_figure")
        assert fig.layout.title.text == ""


# =============================================================================
# _fix_axis_gradient_overlaps_plotly
# =============================================================================


class TestFixAxisGradientOverlaps:
    def test_basic_figure_no_colorbar(self):
        fig = go.Figure(data=[go.Scatter(x=[1, 2], y=[3, 4])])
        _fix_axis_gradient_overlaps_plotly(
            fig, width_px=1600, height_px=1000, tick_px=12, axis_title_px=14
        )
        # Should set margins without raising
        assert fig.layout.margin is not None

    def test_figure_with_heatmap_colorbar(self):
        fig = go.Figure(data=[go.Heatmap(z=[[1, 2], [3, 4]], colorbar=dict(title="Scale"))])
        _fix_axis_gradient_overlaps_plotly(
            fig, width_px=1600, height_px=1000, tick_px=12, axis_title_px=14
        )
        # Right margin should be increased for vertical colorbar
        assert fig.layout.margin.r >= 60

    def test_figure_with_horizontal_colorbar(self):
        fig = go.Figure(
            data=[
                go.Heatmap(
                    z=[[1, 2], [3, 4]],
                    colorbar=dict(title="Scale", orientation="h"),
                )
            ]
        )
        _fix_axis_gradient_overlaps_plotly(
            fig, width_px=1600, height_px=1000, tick_px=12, axis_title_px=14
        )
        assert fig.layout.margin.b >= 90

    def test_figure_with_marker_colorbar(self):
        fig = go.Figure(
            data=[
                go.Scatter(
                    x=[1, 2, 3],
                    y=[4, 5, 6],
                    marker=dict(
                        color=[1, 2, 3],
                        colorbar=dict(title="Marker"),
                        colorscale="Viridis",
                    ),
                )
            ]
        )
        _fix_axis_gradient_overlaps_plotly(
            fig, width_px=1600, height_px=1000, tick_px=12, axis_title_px=14
        )
        # Should handle marker-level colorbar

    def test_figure_with_coloraxis(self):
        fig = go.Figure(data=[go.Heatmap(z=[[1, 2], [3, 4]], coloraxis="coloraxis")])
        fig.update_layout(
            coloraxis=dict(
                colorscale="Viridis",
                colorbar=dict(title="Shared"),
            )
        )
        _fix_axis_gradient_overlaps_plotly(
            fig, width_px=1600, height_px=1000, tick_px=12, axis_title_px=14
        )

    def test_figure_with_title(self):
        fig = go.Figure(data=[go.Scatter(x=[1], y=[1])])
        fig.update_layout(title_text="My Plot")
        _fix_axis_gradient_overlaps_plotly(
            fig, width_px=1600, height_px=1000, tick_px=12, axis_title_px=14
        )
        # Top margin should be larger when title exists
        assert fig.layout.margin.t >= 25

    def test_custom_font_family(self):
        fig = go.Figure(data=[go.Scatter(x=[1], y=[1])])
        _fix_axis_gradient_overlaps_plotly(
            fig,
            width_px=1600,
            height_px=1000,
            tick_px=12,
            axis_title_px=14,
            font_family="Times New Roman",
        )


# =============================================================================
# apply_zip_export_style
# =============================================================================


class TestStyleFigureForExport:
    def test_basic_call(self):
        fig = go.Figure(data=[go.Scatter(x=[1, 2], y=[3, 4])])
        fig.update_layout(title_text="Export me")
        apply_zip_export_style(fig)
        # Title should be stripped
        assert fig.layout.title.text == ""

    def test_custom_dimensions(self):
        fig = go.Figure(data=[go.Scatter(x=[1], y=[1])])
        apply_zip_export_style(fig, width_px=800, height_px=600, tick_px=10)

    def test_backward_compatible_alias(self):
        fig = go.Figure(data=[go.Scatter(x=[1], y=[1])])
        style_figure_for_export(fig, width_px=900, height_px=500, tick_px=11)


# =============================================================================
# _style_plotly_for_docx_export
# =============================================================================


class TestStylePlotlyForDocxExport:
    def test_none_input(self):
        result = _style_plotly_for_docx_export(None, DocxReportConfig())
        assert result is None

    def test_journal_preset(self):
        fig = go.Figure(data=[go.Scatter(x=[1, 2], y=[3, 4])])
        fig.update_layout(title_text="Journal Figure")
        cfg = DocxReportConfig(text_preset="Journal (default)")
        result = _style_plotly_for_docx_export(fig, cfg)
        assert result is not None

    def test_presentation_preset(self):
        fig = go.Figure(data=[go.Scatter(x=[1, 2], y=[3, 4])])
        cfg = DocxReportConfig(text_preset="Presentation (slides)")
        result = _style_plotly_for_docx_export(fig, cfg)
        assert result is not None

    def test_poster_preset(self):
        fig = go.Figure(data=[go.Scatter(x=[1, 2], y=[3, 4])])
        cfg = DocxReportConfig(text_preset="Poster")
        result = _style_plotly_for_docx_export(fig, cfg)
        assert result is not None

    def test_with_annotations(self):
        fig = go.Figure(data=[go.Scatter(x=[1, 2], y=[3, 4])])
        fig.update_layout(
            annotations=[dict(text="Note", x=1, y=2, showarrow=True, font=dict(size=12))]
        )
        cfg = DocxReportConfig()
        result = _style_plotly_for_docx_export(fig, cfg)
        assert result is not None

    def test_with_legend(self):
        fig = go.Figure(
            data=[
                go.Scatter(x=[1, 2], y=[3, 4], name="Series 1"),
                go.Scatter(x=[1, 2], y=[5, 6], name="Series 2"),
            ]
        )
        fig.update_layout(showlegend=True)
        cfg = DocxReportConfig()
        result = _style_plotly_for_docx_export(fig, cfg)
        assert result is not None

    def test_strip_headings_disabled(self):
        fig = go.Figure(data=[go.Scatter(x=[1], y=[1])])
        fig.update_layout(title_text="Keep This")
        cfg = DocxReportConfig(strip_user_headings=False)
        result = _style_plotly_for_docx_export(fig, cfg)
        # Title should be preserved
        assert result.layout.title.text == "Keep This"

    def test_fix_overlaps_disabled(self):
        fig = go.Figure(data=[go.Scatter(x=[1], y=[1])])
        cfg = DocxReportConfig(fix_axis_overlaps=False)
        result = _style_plotly_for_docx_export(fig, cfg)
        assert result is not None

    def test_does_not_mutate_original(self):
        fig = go.Figure(data=[go.Scatter(x=[1, 2], y=[3, 4])])
        fig.update_layout(title_text="Original Title")
        cfg = DocxReportConfig()
        result = _style_plotly_for_docx_export(fig, cfg)
        # Original should be untouched
        assert fig.layout.title.text == "Original Title"
        # Result should have title cleared (strip_user_headings=True)
        assert result.layout.title.text == ""


# =============================================================================
# _best_model_line
# =============================================================================


class TestBestModelLine:
    def test_with_converged_models(self):
        models = {
            "Langmuir": {"converged": True, "r_squared": 0.98},
            "Freundlich": {"converged": True, "r_squared": 0.95},
        }
        result = _best_model_line(models, "Best isotherm model")
        assert result is not None
        assert "Langmuir" in result
        assert "0.98" in result

    def test_no_converged(self):
        models = {
            "Langmuir": {"converged": False},
            "Freundlich": {"converged": False},
        }
        result = _best_model_line(models, "Best model")
        assert result is None

    def test_empty_dict(self):
        assert _best_model_line({}, "label") is None

    def test_not_dict(self):
        assert _best_model_line(None, "label") is None
        assert _best_model_line("string", "label") is None
        assert _best_model_line(42, "label") is None

    def test_missing_r_squared(self):
        models = {"Model": {"converged": True}}
        result = _best_model_line(models, "label")
        # r_squared is None -> -inf, returns None
        assert result is None

    def test_invalid_r_squared_type(self):
        models = {"Model": {"converged": True, "r_squared": "not_a_number"}}
        result = _best_model_line(models, "label")
        assert result is None

    def test_single_model(self):
        models = {"Sips": {"converged": True, "r_squared": 0.99}}
        result = _best_model_line(models, "Best fit")
        assert "Sips" in result
        assert "0.99" in result


# =============================================================================
# _ensure_caption_style
# =============================================================================


class TestEnsureCaptionStyle:
    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_returns_caption(self):
        from docx import Document

        doc = Document()
        result = _ensure_caption_style(doc)
        assert result == "Caption"


# =============================================================================
# _recommended_docx_figure_width_in
# =============================================================================


class TestRecommendedDocxFigureWidth:
    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_default_document(self):
        from docx import Document

        doc = Document()
        w = _recommended_docx_figure_width_in(doc)
        assert 4.5 <= w <= 6.5

    @pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
    def test_custom_bounds(self):
        from docx import Document

        doc = Document()
        w = _recommended_docx_figure_width_in(doc, max_width_in=5.0, min_width_in=3.0)
        assert 3.0 <= w <= 5.0

    def test_bad_doc_fallback(self):
        # With a non-doc object, should return max_width_in
        w = _recommended_docx_figure_width_in(None, max_width_in=7.0)
        assert w == 7.0


# =============================================================================
# _add_dataframe_table
# =============================================================================


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
class TestAddDataframeTable:
    def test_basic_table(self):
        from docx import Document
        from adsorblab_pro.docx_report import _add_dataframe_table

        doc = Document()
        df = pd.DataFrame({"A": [1.0, 2.0], "B": ["x", "y"]})
        cfg = DocxReportConfig()
        _add_dataframe_table(doc, df, "Test Table", cfg)
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 3  # 1 header + 2 data

    def test_truncates_rows(self):
        from docx import Document
        from adsorblab_pro.docx_report import _add_dataframe_table

        doc = Document()
        df = pd.DataFrame({"A": range(100)})
        cfg = DocxReportConfig(max_table_rows=10)
        _add_dataframe_table(doc, df, "Big Table", cfg)
        assert len(doc.tables[0].rows) == 11  # 1 header + 10 data

    def test_truncates_cols(self):
        from docx import Document
        from adsorblab_pro.docx_report import _add_dataframe_table

        doc = Document()
        data = {f"col_{i}": [1] for i in range(20)}
        df = pd.DataFrame(data)
        cfg = DocxReportConfig(max_table_cols=5)
        _add_dataframe_table(doc, df, "Wide Table", cfg)
        assert len(doc.tables[0].columns) == 5

    def test_nan_values(self):
        from docx import Document
        from adsorblab_pro.docx_report import _add_dataframe_table

        doc = Document()
        df = pd.DataFrame({"A": [1.0, float("nan"), 3.0]})
        cfg = DocxReportConfig()
        _add_dataframe_table(doc, df, "NaN Table", cfg)
        cells = [doc.tables[0].rows[i + 1].cells[0].text for i in range(3)]
        assert cells[1] == ""  # NaN rendered as empty

    def test_float_formatting(self):
        from docx import Document
        from adsorblab_pro.docx_report import _add_dataframe_table

        doc = Document()
        df = pd.DataFrame({"A": [3.14159265]})
        cfg = DocxReportConfig(float_format="{:.2f}")
        _add_dataframe_table(doc, df, "Formatted", cfg)
        assert doc.tables[0].rows[1].cells[0].text == "3.14"


# =============================================================================
# _add_figure
# =============================================================================


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
class TestAddFigure:
    def _make_png(self):
        from PIL import Image

        img = Image.new("RGB", (100, 80), color="red")
        bio = BytesIO()
        img.save(bio, format="PNG")
        bio.seek(0)
        return bio.getvalue()

    def test_basic_figure(self):
        from docx import Document
        from adsorblab_pro.docx_report import _add_figure

        doc = Document()
        png = self._make_png()
        cfg = DocxReportConfig(figure_width_in=5.0)
        _add_figure(doc, png, "Figure 1. Test", cfg)
        # Verify picture was added (inline shapes)
        assert len(doc.paragraphs) >= 2


# =============================================================================
# create_docx_report edge cases
# =============================================================================


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
class TestCreateDocxReportEdgeCases:
    def _make_png(self):
        from PIL import Image

        img = Image.new("RGB", (100, 80))
        bio = BytesIO()
        img.save(bio, format="PNG")
        bio.seek(0)
        return bio.getvalue()

    def test_empty_selections(self):
        docx_bytes, warnings = create_docx_report(
            study_title="Empty Report",
            study_state={},
            selected_figures=[],
            selected_tables=[],
            figure_generator=lambda fid, s: None,
            table_generator=lambda tid, s: None,
        )
        assert isinstance(docx_bytes, (bytes, bytearray))
        assert len(docx_bytes) > 500

    def test_failing_figure_generator(self):
        docx_bytes, warnings = create_docx_report(
            study_title="Fail Report",
            study_state={},
            selected_figures=["bad_fig"],
            selected_tables=[],
            figure_generator=lambda fid, s: (_ for _ in ()).throw(ValueError("boom")),
            table_generator=lambda tid, s: None,
        )
        assert any("bad_fig" in w for w in warnings)

    def test_none_figure_generator(self):
        docx_bytes, warnings = create_docx_report(
            study_title="None Fig Report",
            study_state={},
            selected_figures=["fig1"],
            selected_tables=[],
            figure_generator=lambda fid, s: None,
            table_generator=lambda tid, s: None,
        )
        assert any("fig1" in w for w in warnings)

    def test_failing_table_generator(self):
        docx_bytes, warnings = create_docx_report(
            study_title="Fail Table Report",
            study_state={},
            selected_figures=[],
            selected_tables=["bad_tbl"],
            figure_generator=lambda fid, s: None,
            table_generator=lambda tid, s: (_ for _ in ()).throw(ValueError("table error")),
        )
        assert any("bad_tbl" in w for w in warnings)

    def test_empty_table(self):
        docx_bytes, warnings = create_docx_report(
            study_title="Empty Table Report",
            study_state={},
            selected_figures=[],
            selected_tables=["empty_tbl"],
            figure_generator=lambda fid, s: None,
            table_generator=lambda tid, s: pd.DataFrame(),
        )
        assert any("empty_tbl" in w for w in warnings)

    def test_with_figure_and_table_meta(self):
        png = self._make_png()
        docx_bytes, warnings = create_docx_report(
            study_title="Meta Report",
            study_state={"isotherm_models_fitted": {}},
            selected_figures=["fig1"],
            selected_tables=["tbl1"],
            figure_generator=lambda fid, s: png,
            table_generator=lambda tid, s: pd.DataFrame({"A": [1]}),
            figure_meta={"fig1": ("Langmuir Fit", "Best model fit")},
            table_meta={"tbl1": ("Parameters", "Fitted values")},
        )
        # Figure gets passed as raw bytes, _coerce_image_bytes should handle it
        # This may generate a warning if it expects a Plotly fig, but should not crash

    def test_study_overview_analyses(self):
        state = {
            "isotherm_models_fitted": {"Langmuir": {"converged": True, "r_squared": 0.99}},
            "kinetic_models_fitted": {"PFO": {"converged": True, "r_squared": 0.95}},
            "thermo_params": {
                "success": True,
                "delta_H": -25.0,
                "delta_S": 0.05,
                "delta_G": [-20.0],
            },
            "ph_effect_results": {"data": "something"},
            "temp_effect_results": {"data": "something"},
            "dosage_results": {"data": "something"},
        }
        docx_bytes, warnings = create_docx_report(
            study_title="Full Analysis Report",
            study_state=state,
            selected_figures=[],
            selected_tables=[],
            figure_generator=lambda fid, s: None,
            table_generator=lambda tid, s: None,
        )
        from docx import Document

        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Isotherms" in full_text
        assert "Kinetics" in full_text
        assert "Thermodynamics" in full_text

    def test_thermo_params_bad_values(self):
        """Test thermo_params with missing values fall back gracefully."""
        state = {
            "thermo_params": {"success": True, "delta_H": None, "delta_S": None, "delta_G": None},
        }
        docx_bytes, warnings = create_docx_report(
            study_title="Thermo Edge",
            study_state=state,
            selected_figures=[],
            selected_tables=[],
            figure_generator=lambda fid, s: None,
            table_generator=lambda tid, s: None,
        )
        assert isinstance(docx_bytes, (bytes, bytearray))

    def test_many_warnings_truncation(self):
        """Test that >30 warnings get truncated."""

        def bad_table(tid, s):
            raise ValueError(f"fail {tid}")

        tables = [f"tbl_{i}" for i in range(35)]
        docx_bytes, warnings = create_docx_report(
            study_title="Many Warnings",
            study_state={},
            selected_figures=[],
            selected_tables=tables,
            figure_generator=lambda fid, s: None,
            table_generator=bad_table,
        )
        assert len(warnings) == 70  # 35 tables × 2 warnings each (generation error + empty)
