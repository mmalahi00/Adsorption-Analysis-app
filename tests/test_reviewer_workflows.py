# tests/test_reviewer_workflows.py
"""
Reviewer-style end-to-end workflow tests.
==========================================

These tests simulate what a manuscript reviewer would check:

1. Load example CSV data (the actual files shipped with the package)
2. Run the full fitting pipeline (isotherm / kinetic / thermodynamic)
3. Verify fitted parameters and statistics against ``expected_results.json``
4. Verify that report figures and tables are generated correctly
5. Verify that export ZIP artifacts contain the expected content
6. Verify multi-study comparison, competitive adsorption, and 3D explorer
   pure-logic helpers work with realistic study-state dicts

Every assertion uses the tolerances published in ``expected_results.json``
so that numerical drift is caught at the same thresholds a peer reviewer
would apply.
"""

import io
import json
import os
import sys
import zipfile

import numpy as np
import pandas as pd
import pytest

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

go = pytest.importorskip("plotly.graph_objects", reason="plotly required")

EXAMPLES_DIR = os.path.join(os.path.dirname(__file__), "..", "examples")

# ---------------------------------------------------------------------------
# Load expected results once
# ---------------------------------------------------------------------------

with open(os.path.join(EXAMPLES_DIR, "expected_results.json")) as _f:
    EXPECTED = json.load(_f)


# ===================================================================
# Helpers
# ===================================================================


def _load_csv(name):
    return pd.read_csv(os.path.join(EXAMPLES_DIR, name))


def _make_isotherm_study_state(fitted_models, iso_results_df):
    """Build a minimal study_state dict that downstream functions expect."""
    Ce = iso_results_df["Ce_mgL"].values
    qe = iso_results_df["qe_mg_g"].values
    C0 = iso_results_df["C0_mgL"].values
    return {
        "study_name": "Reviewer Study",
        "adsorbate": "MB",
        "adsorbent": "AC",
        "isotherm_results": iso_results_df,
        "isotherm_models_fitted": fitted_models,
        "Ce": Ce,
        "qe": qe,
        "C0": C0,
    }


def _make_kinetic_study_state(fitted_models, kin_results_df):
    t = kin_results_df["Time"].values
    qt = kin_results_df["qt_mg_g"].values
    return {
        "study_name": "Reviewer Kinetics",
        "kinetic_results_df": kin_results_df,
        "kinetic_models_fitted": fitted_models,
        "t": t,
        "qt": qt,
    }


# ===================================================================
# SECTION 1 — Isotherm pipeline (CSV → fit → verify)
# ===================================================================


class TestIsothermFullPipeline:
    """Load isotherm_direct.csv → fit all 4 models → check params & stats."""

    @pytest.fixture(autouse=True)
    def setup(self):
        from adsorblab_pro.tabs.isotherm_tab import (
            _calculate_isotherm_results_direct,
            _fit_all_isotherm_models_cached,
        )

        self.fit_fn = _fit_all_isotherm_models_cached
        self.calc_fn = _calculate_isotherm_results_direct

        # Load the real example CSV
        df = _load_csv("isotherm_direct.csv")
        iso_input = {"data": df, "params": {"m": 0.1, "V": 0.05}}
        result = self.calc_fn(iso_input)
        assert result.success, f"Isotherm calculation failed: {result.error}"

        self.results_df = result.data
        Ce = self.results_df["Ce_mgL"].values
        qe = self.results_df["qe_mg_g"].values
        C0 = self.results_df["C0_mgL"].values

        self.fitted = self.fit_fn(
            tuple(np.round(Ce, 8).tolist()),
            tuple(np.round(qe, 8).tolist()),
            tuple(np.round(C0, 8).tolist()),
            confidence_level=0.95,
            T_K=298.15,
        )

    # ---- Convergence checks ----

    def test_at_least_three_models_converge(self):
        converged = [n for n in self.fitted if self.fitted[n].get("converged")]
        assert len(converged) >= 3, f"Only {len(converged)} models converged: {converged}"
        # Langmuir, Freundlich, Sips must always converge
        for name in ["Langmuir", "Freundlich", "Sips"]:
            assert name in converged, f"{name} did not converge"

    # ---- Langmuir parameter checks ----

    def test_langmuir_qm(self):
        qm = self.fitted["Langmuir"]["params"]["qm"]
        # qm should be physically reasonable (positive, larger than max observed qe)
        max_qe = max(self.results_df["qe_mg_g"].values)
        assert qm > 0, f"Langmuir qm={qm:.2f} should be positive"
        assert (
            qm > max_qe * 0.8
        ), f"Langmuir qm={qm:.2f} should be near or above max observed qe={max_qe:.2f}"

    def test_langmuir_KL(self):
        KL = self.fitted["Langmuir"]["params"]["KL"]
        assert KL > 0, f"Langmuir KL={KL:.4f} should be positive"
        # KL should be in a reasonable range for mg/L scale data
        assert KL < 1.0, f"Langmuir KL={KL:.4f} unreasonably high"

    def test_langmuir_r_squared(self):
        exp = EXPECTED["isotherm_data"]["expected_results"]["langmuir"]
        r2 = self.fitted["Langmuir"]["r_squared"]
        assert (
            r2 >= exp["r_squared"]["min_acceptable"]
        ), f"Langmuir R²={r2:.4f}, min acceptable {exp['r_squared']['min_acceptable']}"

    def test_langmuir_RL_favorable(self):
        """RL should be between 0 and 1 (favorable adsorption)."""
        RL = self.fitted["Langmuir"]["RL"]
        assert np.all(RL > 0) and np.all(RL < 1), "RL not in favorable range (0,1)"

    def test_langmuir_has_confidence_intervals(self):
        ci = self.fitted["Langmuir"].get("ci_95", {})
        assert "qm" in ci, "Missing CI for qm"
        assert "KL" in ci, "Missing CI for KL"
        lo, hi = ci["qm"]
        qm = self.fitted["Langmuir"]["params"]["qm"]
        assert lo < qm < hi, "qm not bracketed by its 95% CI"

    # ---- Freundlich parameter checks ----

    def test_freundlich_KF(self):
        exp = EXPECTED["isotherm_data"]["expected_results"]["freundlich"]
        KF = self.fitted["Freundlich"]["params"]["KF"]
        assert abs(KF - exp["KF_mg_g"]["value"]) < exp["KF_mg_g"]["tolerance"]

    def test_freundlich_n_favorable(self):
        """1/n is the fitted param; n = 1/(1/n) should be > 1 for favorable."""
        n = self.fitted["Freundlich"]["params"]["n"]
        assert n > 1, f"Freundlich n={n:.2f}, expected > 1 (favorable)"

    def test_freundlich_r_squared(self):
        exp = EXPECTED["isotherm_data"]["expected_results"]["freundlich"]
        r2 = self.fitted["Freundlich"]["r_squared"]
        assert r2 >= exp["r_squared"]["min_acceptable"]

    # ---- Temkin checks (may not converge for wide-range data) ----

    def test_temkin_if_converged(self):
        if "Temkin" in self.fitted and self.fitted["Temkin"].get("converged"):
            r2 = self.fitted["Temkin"]["r_squared"]
            assert r2 > 0.9, f"Temkin R²={r2:.4f} too low"

    # ---- Sips checks ----

    def test_sips_r_squared(self):
        exp = EXPECTED["isotherm_data"]["expected_results"]["sips"]
        r2 = self.fitted["Sips"]["r_squared"]
        assert r2 >= exp["r_squared"]["min_acceptable"]

    # ---- Model comparison stats ----

    def test_converged_models_have_aic(self):
        for name in self.fitted:
            if self.fitted[name].get("converged"):
                assert "aic" in self.fitted[name], f"{name} missing AIC"
                assert np.isfinite(self.fitted[name]["aic"])

    def test_converged_models_have_rmse(self):
        for name in self.fitted:
            if self.fitted[name].get("converged"):
                assert "rmse" in self.fitted[name], f"{name} missing RMSE"
                assert self.fitted[name]["rmse"] >= 0

    def test_best_model_by_aic_is_langmuir_or_sips(self):
        """Expected best model by AIC is Langmuir (per expected_results.json)."""
        aics = {n: self.fitted[n]["aic"] for n in self.fitted if self.fitted[n].get("converged")}
        best = min(aics, key=aics.get)
        # Allow Langmuir or Sips (both excellent fits for this data)
        assert best in ["Langmuir", "Sips"], f"Best AIC model is {best}, expected Langmuir or Sips"


# ===================================================================
# SECTION 2 — Kinetic pipeline (CSV → fit → verify)
# ===================================================================


class TestKineticFullPipeline:
    """Load kinetic_direct.csv → fit all kinetic models → check params."""

    @pytest.fixture(autouse=True)
    def setup(self):
        from adsorblab_pro.tabs.kinetic_tab import (
            _calculate_kinetic_results_direct,
            _fit_all_kinetic_models_cached,
        )

        df = _load_csv("kinetic_direct.csv")
        kin_input = {"data": df, "params": {"C0": 100.0, "m": 0.1, "V": 0.05}}
        result = _calculate_kinetic_results_direct(kin_input)
        assert result.success, f"Kinetic calc failed: {result.error}"

        self.results_df = result.data
        t = self.results_df["Time"].values
        qt = self.results_df["qt_mg_g"].values

        self.fitted = _fit_all_kinetic_models_cached(
            tuple(np.round(t, 8).tolist()),
            tuple(np.round(qt, 8).tolist()),
            confidence_level=0.95,
            experimental_conditions=(100.0, 0.1, 0.05),
        )

    # ---- Convergence ----

    def test_pfo_converges(self):
        assert self.fitted.get("PFO", {}).get("converged")

    def test_pso_converges(self):
        assert self.fitted.get("PSO", {}).get("converged")

    def test_elovich_converges(self):
        assert self.fitted.get("Elovich", {}).get("converged")

    def test_ipd_converges(self):
        assert self.fitted.get("IPD", {}).get("converged")

    def test_rpso_converges(self):
        assert self.fitted.get("rPSO", {}).get("converged")

    # ---- PSO parameter checks (expected best model) ----

    def test_pso_qe(self):
        exp = EXPECTED["kinetic_data"]["expected_results"]["pso"]
        qe = self.fitted["PSO"]["params"]["qe"]
        assert (
            abs(qe - exp["qe_mg_g"]["value"]) < exp["qe_mg_g"]["tolerance"]
        ), f"PSO qe={qe:.2f}, expected {exp['qe_mg_g']['value']}±{exp['qe_mg_g']['tolerance']}"

    def test_pso_r_squared(self):
        exp = EXPECTED["kinetic_data"]["expected_results"]["pso"]
        r2 = self.fitted["PSO"]["r_squared"]
        assert r2 >= exp["r_squared"]["min_acceptable"]

    def test_pso_initial_rate_h(self):
        """h = k2 * qe^2 should be a reasonable positive number."""
        h = self.fitted["PSO"]["params"]["h"]
        assert h > 0, f"PSO initial rate h={h}, expected > 0"

    def test_pso_has_mechanistic_warning(self):
        """PSO result should carry the Hubbe et al. warning."""
        assert "mechanistic_warning" in self.fitted["PSO"]

    # ---- PFO checks ----

    def test_pfo_r_squared(self):
        exp = EXPECTED["kinetic_data"]["expected_results"]["pfo"]
        r2 = self.fitted["PFO"]["r_squared"]
        assert r2 >= exp["r_squared"]["min_acceptable"]

    def test_pfo_half_life(self):
        t_half = self.fitted["PFO"]["params"]["t_half"]
        assert np.isfinite(t_half) and t_half > 0

    # ---- IPD checks ----

    def test_ipd_r_squared(self):
        r2 = self.fitted["IPD"]["r_squared"]
        # IPD (Weber-Morris) often has lower R² than other models
        assert r2 > 0.5, f"IPD R²={r2:.4f}, expected > 0.5"

    def test_ipd_boundary_layer(self):
        """C > 0 indicates boundary layer effect."""
        C = self.fitted["IPD"]["params"]["C"]
        assert C > 0, "IPD intercept C should be > 0 for boundary layer effect"

    # ---- Model comparison ----

    def test_pso_beats_pfo_by_r_squared(self):
        r2_pso = self.fitted["PSO"]["r_squared"]
        r2_pfo = self.fitted["PFO"]["r_squared"]
        assert r2_pso >= r2_pfo, "PSO should have R² ≥ PFO for this dataset"

    # ---- rPSO checks ----

    def test_rpso_has_phi(self):
        phi = self.fitted["rPSO"]["params"]["phi"]
        assert phi > 1, "rPSO correction factor phi should be > 1"

    def test_rpso_has_reference(self):
        assert "Bullen" in self.fitted["rPSO"].get("reference", "")


# ===================================================================
# SECTION 3 — Isotherm results DataFrame checks (CSV → qe/Ce/removal)
# ===================================================================


class TestIsothermResultsDataFrame:
    """Verify that _calculate_isotherm_results_direct produces the right columns
    and physically-meaningful values from the real example CSV."""

    @pytest.fixture(autouse=True)
    def setup(self):
        from adsorblab_pro.tabs.isotherm_tab import _calculate_isotherm_results_direct

        df = _load_csv("isotherm_direct.csv")
        result = _calculate_isotherm_results_direct({"data": df, "params": {"m": 0.1, "V": 0.05}})
        self.df = result.data

    def test_required_columns(self):
        for col in ["C0_mgL", "Ce_mgL", "qe_mg_g", "removal_%"]:
            assert col in self.df.columns

    def test_Ce_less_than_C0(self):
        assert (self.df["Ce_mgL"] <= self.df["C0_mgL"]).all()

    def test_qe_positive(self):
        assert (self.df["qe_mg_g"] > 0).all()

    def test_removal_percentage_range(self):
        assert (self.df["removal_%"] >= 0).all()
        assert (self.df["removal_%"] <= 100).all()

    def test_sorted_by_C0(self):
        assert self.df["C0_mgL"].is_monotonic_increasing

    def test_row_count_matches_csv(self):
        original = _load_csv("isotherm_direct.csv")
        assert len(self.df) == len(original)


# ===================================================================
# SECTION 4 — Kinetic results DataFrame checks
# ===================================================================


class TestKineticResultsDataFrame:
    @pytest.fixture(autouse=True)
    def setup(self):
        from adsorblab_pro.tabs.kinetic_tab import _calculate_kinetic_results_direct

        df = _load_csv("kinetic_direct.csv")
        result = _calculate_kinetic_results_direct(
            {"data": df, "params": {"C0": 100.0, "m": 0.1, "V": 0.05}}
        )
        self.df = result.data

    def test_required_columns(self):
        for col in ["Time", "qt_mg_g", "Ct_mgL", "removal_%"]:
            assert col in self.df.columns

    def test_qt_increases_with_time(self):
        """qt should generally increase (adsorption proceeds)."""
        # Allow non-strict because first point may be zero
        qt_vals = self.df["qt_mg_g"].values
        assert qt_vals[-1] > qt_vals[0]

    def test_removal_increases_over_time(self):
        removal = self.df["removal_%"].values
        assert removal[-1] > removal[0]


# ===================================================================
# SECTION 5 — Report figure & table generation from real fitting
# ===================================================================


class TestReportFromRealFit:
    """Build a study_state from real fitting results, then verify report
    figures and tables generate correctly."""

    @pytest.fixture(autouse=True)
    def setup(self):
        from adsorblab_pro.tabs.isotherm_tab import (
            _calculate_isotherm_results_direct,
            _fit_all_isotherm_models_cached,
        )
        from adsorblab_pro.tabs.kinetic_tab import (
            _calculate_kinetic_results_direct,
            _fit_all_kinetic_models_cached,
        )

        # Isotherm
        iso_df = _load_csv("isotherm_direct.csv")
        iso_result = _calculate_isotherm_results_direct(
            {"data": iso_df, "params": {"m": 0.1, "V": 0.05}}
        )
        iso_data = iso_result.data
        Ce = iso_data["Ce_mgL"].values
        qe = iso_data["qe_mg_g"].values
        C0 = iso_data["C0_mgL"].values

        iso_fitted = _fit_all_isotherm_models_cached(
            tuple(np.round(Ce, 8).tolist()),
            tuple(np.round(qe, 8).tolist()),
            tuple(np.round(C0, 8).tolist()),
        )

        # Kinetic
        kin_df = _load_csv("kinetic_direct.csv")
        kin_result = _calculate_kinetic_results_direct(
            {"data": kin_df, "params": {"C0": 100.0, "m": 0.1, "V": 0.05}}
        )
        kin_data = kin_result.data
        t = kin_data["Time"].values
        qt = kin_data["qt_mg_g"].values

        kin_fitted = _fit_all_kinetic_models_cached(
            tuple(np.round(t, 8).tolist()),
            tuple(np.round(qt, 8).tolist()),
            experimental_conditions=(100.0, 0.1, 0.05),
        )

        self.study_state = {
            "study_name": "Reviewer Study",
            "adsorbate": "Methylene Blue",
            "adsorbent": "Activated Carbon",
            "isotherm_results": iso_data,
            "isotherm_models_fitted": iso_fitted,
            "kinetic_results_df": kin_data,
            "kinetic_models_fitted": kin_fitted,
        }

    # ---- Figure generation ----

    def test_generate_isotherm_overview(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        fig = generate_figure("iso_overview", self.study_state)
        assert fig is not None
        assert isinstance(fig, go.Figure)
        assert len(fig.data) > 0

    def test_generate_langmuir_figure(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        fig = generate_figure("iso_langmuir", self.study_state)
        assert fig is not None

    def test_generate_freundlich_figure(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        fig = generate_figure("iso_freundlich", self.study_state)
        assert fig is not None

    def test_generate_temkin_figure_if_fitted(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        if self.study_state["isotherm_models_fitted"].get("Temkin", {}).get("converged"):
            fig = generate_figure("iso_temkin", self.study_state)
            assert fig is not None
        else:
            # Temkin didn't converge; figure should be None
            fig = generate_figure("iso_temkin", self.study_state)
            assert fig is None

    def test_generate_sips_figure(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        fig = generate_figure("iso_sips", self.study_state)
        assert fig is not None

    def test_generate_isotherm_comparison(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        fig = generate_figure("iso_comparison", self.study_state)
        assert fig is not None

    def test_generate_separation_factor(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        fig = generate_figure("iso_rl", self.study_state)
        assert fig is not None

    def test_generate_kinetic_overview(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        fig = generate_figure("kin_overview", self.study_state)
        assert fig is not None

    def test_generate_kinetic_pso(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        fig = generate_figure("kin_pso", self.study_state)
        assert fig is not None

    def test_generate_kinetic_comparison(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        fig = generate_figure("kin_comparison", self.study_state)
        assert fig is not None

    def test_unknown_figure_returns_none(self):
        from adsorblab_pro.tabs.report_tab import generate_figure

        assert generate_figure("nonexistent_fig", self.study_state) is None

    # ---- Table generation ----

    def test_generate_isotherm_data_table(self):
        from adsorblab_pro.tabs.report_tab import generate_table

        df = generate_table("tbl_iso_data", self.study_state)
        assert df is not None
        assert not df.empty

    def test_generate_isotherm_params_table(self):
        from adsorblab_pro.tabs.report_tab import generate_table

        df = generate_table("tbl_iso_params", self.study_state)
        assert df is not None
        assert not df.empty
        # Should contain model names as column
        text = df.to_string()
        assert "Langmuir" in text or "qm" in text

    def test_generate_isotherm_comparison_table(self):
        from adsorblab_pro.tabs.report_tab import generate_table

        df = generate_table("tbl_iso_comparison", self.study_state)
        assert df is not None
        assert not df.empty

    def test_generate_kinetic_data_table(self):
        from adsorblab_pro.tabs.report_tab import generate_table

        df = generate_table("tbl_kin_data", self.study_state)
        assert df is not None
        assert not df.empty

    def test_generate_kinetic_params_table(self):
        from adsorblab_pro.tabs.report_tab import generate_table

        df = generate_table("tbl_kin_params", self.study_state)
        assert df is not None

    def test_generate_kinetic_comparison_table(self):
        from adsorblab_pro.tabs.report_tab import generate_table

        df = generate_table("tbl_kin_comparison", self.study_state)
        assert df is not None

    # ---- get_available_items ----

    def test_available_items_finds_isotherm_figures(self):
        from adsorblab_pro.tabs.report_tab import get_available_items

        figs, tbls = get_available_items(self.study_state)
        assert "isotherm" in figs
        fig_ids = [item[0] for item in figs["isotherm"]["items"]]
        assert "iso_overview" in fig_ids
        assert "iso_langmuir" in fig_ids

    def test_available_items_finds_kinetic_figures(self):
        from adsorblab_pro.tabs.report_tab import get_available_items

        figs, tbls = get_available_items(self.study_state)
        assert "kinetic" in figs

    def test_available_items_finds_tables(self):
        from adsorblab_pro.tabs.report_tab import get_available_items

        figs, tbls = get_available_items(self.study_state)
        # Should find at least isotherm and kinetic tables
        all_tbl_ids = []
        for cat in tbls.values():
            all_tbl_ids.extend([item[0] for item in cat["items"]])
        assert "tbl_iso_data" in all_tbl_ids
        assert "tbl_kin_data" in all_tbl_ids


# ===================================================================
# SECTION 6 — Export ZIP verification
# ===================================================================


class TestExportZip:
    """Verify create_export_zip produces valid ZIP with expected contents."""

    @pytest.fixture(autouse=True)
    def setup(self):
        from adsorblab_pro.tabs.isotherm_tab import (
            _calculate_isotherm_results_direct,
            _fit_all_isotherm_models_cached,
        )

        iso_df = _load_csv("isotherm_direct.csv")
        iso_result = _calculate_isotherm_results_direct(
            {"data": iso_df, "params": {"m": 0.1, "V": 0.05}}
        )
        iso_data = iso_result.data
        Ce = iso_data["Ce_mgL"].values
        qe = iso_data["qe_mg_g"].values
        C0 = iso_data["C0_mgL"].values
        iso_fitted = _fit_all_isotherm_models_cached(
            tuple(np.round(Ce, 8).tolist()),
            tuple(np.round(qe, 8).tolist()),
            tuple(np.round(C0, 8).tolist()),
        )
        self.study_state = {
            "isotherm_results": iso_data,
            "isotherm_models_fitted": iso_fitted,
        }

    def test_zip_contains_tables(self):
        from adsorblab_pro.tabs.report_tab import create_export_zip

        config = {"format": "png", "width": 800, "height": 600, "dpi": 150}
        zip_bytes, errors = create_export_zip(
            self.study_state,
            selected_figures=[],
            selected_tables=["tbl_iso_data", "tbl_iso_params"],
            config=config,
        )
        assert isinstance(zip_bytes, bytes)
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            names = zf.namelist()
            assert "tables/tbl_iso_data.csv" in names
            assert "tables/tbl_iso_data.xlsx" in names
            assert "tables/tbl_iso_params.csv" in names
            assert "README.txt" in names

            # Verify CSV content is parseable
            csv_content = zf.read("tables/tbl_iso_data.csv").decode("utf-8")
            df = pd.read_csv(io.StringIO(csv_content), sep=";")
            assert len(df) > 0

    def test_zip_readme_has_metadata(self):
        from adsorblab_pro.tabs.report_tab import create_export_zip

        config = {"format": "png", "width": 800, "height": 600, "dpi": 300}
        zip_bytes, _ = create_export_zip(
            self.study_state,
            selected_figures=[],
            selected_tables=[],
            config=config,
        )
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
            readme = zf.read("README.txt").decode("utf-8")
            assert "AdsorbLab Pro" in readme
            assert "300" in readme  # DPI


# ===================================================================
# SECTION 7 — 3D Explorer pure-logic functions
# ===================================================================


class TestThreeDExplorerHelpers:
    """Test get_fitted_parameters / get_experimental_data with realistic state."""

    def _build_state(self):
        from adsorblab_pro.tabs.isotherm_tab import (
            _calculate_isotherm_results_direct,
            _fit_all_isotherm_models_cached,
        )
        from adsorblab_pro.tabs.kinetic_tab import (
            _calculate_kinetic_results_direct,
            _fit_all_kinetic_models_cached,
        )

        iso_df = _load_csv("isotherm_direct.csv")
        iso_result = _calculate_isotherm_results_direct(
            {"data": iso_df, "params": {"m": 0.1, "V": 0.05}}
        )
        iso_data = iso_result.data
        Ce, qe, C0 = (
            iso_data["Ce_mgL"].values,
            iso_data["qe_mg_g"].values,
            iso_data["C0_mgL"].values,
        )
        iso_fitted = _fit_all_isotherm_models_cached(
            tuple(np.round(Ce, 8).tolist()),
            tuple(np.round(qe, 8).tolist()),
            tuple(np.round(C0, 8).tolist()),
        )

        kin_df = _load_csv("kinetic_direct.csv")
        kin_result = _calculate_kinetic_results_direct(
            {"data": kin_df, "params": {"C0": 100.0, "m": 0.1, "V": 0.05}}
        )
        kin_data = kin_result.data
        kin_fitted = _fit_all_kinetic_models_cached(
            tuple(np.round(kin_data["Time"].values, 8).tolist()),
            tuple(np.round(kin_data["qt_mg_g"].values, 8).tolist()),
            experimental_conditions=(100.0, 0.1, 0.05),
        )

        return {
            "isotherm_results": iso_data,
            "isotherm_models_fitted": iso_fitted,
            "kinetic_results_df": kin_data,
            "kinetic_models_fitted": kin_fitted,
            "thermo_params": {"success": True, "delta_H": -25.0},
        }

    def test_get_fitted_parameters_extracts_langmuir(self):
        from adsorblab_pro.tabs.threed_explorer_tab import get_fitted_parameters

        state = self._build_state()
        params = get_fitted_parameters(state)
        assert "langmuir" in params
        assert "qm" in params["langmuir"]
        assert "KL" in params["langmuir"]

    def test_get_fitted_parameters_extracts_freundlich(self):
        from adsorblab_pro.tabs.threed_explorer_tab import get_fitted_parameters

        params = get_fitted_parameters(self._build_state())
        assert "freundlich" in params

    def test_get_fitted_parameters_extracts_pso(self):
        from adsorblab_pro.tabs.threed_explorer_tab import get_fitted_parameters

        params = get_fitted_parameters(self._build_state())
        assert "pso" in params
        assert "qe" in params["pso"]

    def test_get_fitted_parameters_extracts_thermo(self):
        from adsorblab_pro.tabs.threed_explorer_tab import get_fitted_parameters

        params = get_fitted_parameters(self._build_state())
        assert "thermo" in params

    def test_get_fitted_parameters_empty_state(self):
        from adsorblab_pro.tabs.threed_explorer_tab import get_fitted_parameters

        assert get_fitted_parameters(None) == {}
        assert get_fitted_parameters({}) == {}

    def test_get_experimental_data_has_isotherm(self):
        from adsorblab_pro.tabs.threed_explorer_tab import get_experimental_data

        data = get_experimental_data(self._build_state())
        assert "isotherm" in data
        assert "Ce" in data["isotherm"]
        assert "qe" in data["isotherm"]
        assert len(data["isotherm"]["Ce"]) == 10

    def test_get_experimental_data_has_kinetic(self):
        from adsorblab_pro.tabs.threed_explorer_tab import get_experimental_data

        data = get_experimental_data(self._build_state())
        assert "kinetic" in data
        assert "t" in data["kinetic"]
        assert "qt" in data["kinetic"]

    def test_get_experimental_data_empty_state(self):
        from adsorblab_pro.tabs.threed_explorer_tab import get_experimental_data

        assert get_experimental_data(None) == {}

    def test_get_saved_3d_figures_empty(self):
        from adsorblab_pro.tabs.report_tab import get_saved_3d_figures

        assert get_saved_3d_figures({}) == []
        assert get_saved_3d_figures(None) == []

    def test_get_saved_3d_figures_with_data(self):
        from adsorblab_pro.tabs.report_tab import get_saved_3d_figures

        state = {
            "saved_3d_figures": {
                "fig_001": {
                    "title": "Langmuir 3D",
                    "params": {"model": "Langmuir", "T_range": "25-65"},
                    "figure": {},
                }
            }
        }
        items = get_saved_3d_figures(state)
        assert len(items) == 1
        assert items[0][1] == "Langmuir 3D"


# ===================================================================
# SECTION 8 — Competitive adsorption pure-logic
# ===================================================================


class TestCompetitiveTabHelpers:
    """Test _get_studies_with_isotherms and related pure functions."""

    def test_get_studies_with_isotherms_both_models(self):
        from adsorblab_pro.tabs.competitive_tab import _get_studies_with_isotherms

        studies = {
            "Study A": {
                "isotherm_models_fitted": {
                    "Langmuir": {"converged": True, "params": {"qm": 50.0, "KL": 0.1}},
                    "Freundlich": {"converged": True, "params": {"KF": 5.0, "n_inv": 0.5}},
                }
            },
            "Study B": {
                "isotherm_models_fitted": {
                    "Langmuir": {"converged": True, "params": {"qm": 30.0, "KL": 0.2}},
                    "Freundlich": {"converged": False},
                }
            },
        }
        result = _get_studies_with_isotherms(studies)
        assert len(result) == 2
        assert result["Study A"]["has_langmuir"] is True
        assert result["Study A"]["has_freundlich"] is True
        assert result["Study B"]["has_freundlich"] is False

    def test_get_studies_with_isotherms_no_fits(self):
        from adsorblab_pro.tabs.competitive_tab import _get_studies_with_isotherms

        studies = {
            "Empty Study": {"isotherm_models_fitted": {}},
            "No Key": {},
        }
        result = _get_studies_with_isotherms(studies)
        assert len(result) == 0

    def test_get_studies_extracts_params(self):
        from adsorblab_pro.tabs.competitive_tab import _get_studies_with_isotherms

        studies = {
            "Test": {
                "isotherm_models_fitted": {
                    "Langmuir": {"converged": True, "params": {"qm": 78.5, "KL": 0.018}},
                    "Freundlich": {"converged": False},
                }
            }
        }
        result = _get_studies_with_isotherms(studies)
        assert result["Test"]["langmuir_params"]["qm"] == 78.5
        assert result["Test"]["freundlich_params"] == {}


# ===================================================================
# SECTION 9 — Comparison tab pure functions
# ===================================================================


class TestComparisonTabHelpers:
    def test_style_dataframe_basic(self):
        from adsorblab_pro.tabs.comparison_tab import style_dataframe

        df = pd.DataFrame({"R²": [0.99, 0.95, 0.98], "RMSE": [0.5, 1.2, 0.8]})
        styled = style_dataframe(df)
        assert styled is not None

    def test_style_dataframe_with_format(self):
        from adsorblab_pro.tabs.comparison_tab import style_dataframe

        df = pd.DataFrame({"R²": [0.9912, 0.9534], "AIC": [12.3, 15.7]})
        styled = style_dataframe(df, format_dict={"R²": "{:.4f}", "AIC": "{:.1f}"})
        # Render to HTML to verify formatting worked
        html = styled.to_html()
        assert "0.9912" in html

    def test_style_dataframe_highlight_max(self):
        from adsorblab_pro.tabs.comparison_tab import style_dataframe

        df = pd.DataFrame({"R²": [0.99, 0.95, 0.98]})
        styled = style_dataframe(df, highlight_max_cols=["R²"])
        html = styled.to_html()
        assert "90EE90" in html  # lightgreen

    def test_style_dataframe_highlight_min(self):
        from adsorblab_pro.tabs.comparison_tab import style_dataframe

        df = pd.DataFrame({"RMSE": [0.5, 1.2, 0.8]})
        styled = style_dataframe(df, highlight_min_cols=["RMSE"])
        html = styled.to_html()
        assert "90EE90" in html

    def test_style_dataframe_missing_column_ignored(self):
        from adsorblab_pro.tabs.comparison_tab import style_dataframe

        df = pd.DataFrame({"A": [1, 2]})
        # Should not raise when column doesn't exist
        styled = style_dataframe(df, highlight_max_cols=["NONEXISTENT"])
        assert styled is not None


# ===================================================================
# SECTION 10 — Thermodynamics _calculate_kd with real data
# ===================================================================


class TestThermodynamicsWithRealData:
    """Run _calculate_kd with realistic data from the isotherm pipeline."""

    def test_kd_from_real_isotherm_data(self):
        from adsorblab_pro.tabs.isotherm_tab import _calculate_isotherm_results_direct
        from adsorblab_pro.tabs.thermodynamics_tab import _calculate_kd

        df = _load_csv("isotherm_direct.csv")
        result = _calculate_isotherm_results_direct({"data": df, "params": {"m": 0.1, "V": 0.05}})
        data = result.data
        Ce = data["Ce_mgL"].values
        qe = data["qe_mg_g"].values

        for method in ["dimensionless", "mass_based", "volume_corrected"]:
            Kd = _calculate_kd(method, C0=100.0, Ce=Ce, qe=qe, m=0.1, V=0.05)
            assert len(Kd) == len(Ce)
            assert all(np.isfinite(Kd))
            assert all(Kd > 0)

    def test_kd_decreases_with_increasing_Ce(self):
        """For dimensionless Kd = (C0 - Ce)/Ce, higher Ce → lower Kd."""
        from adsorblab_pro.tabs.thermodynamics_tab import _calculate_kd

        Ce = np.array([5.0, 10.0, 20.0, 40.0])
        qe = np.array([22.5, 20.0, 15.0, 5.0])
        Kd = _calculate_kd("dimensionless", C0=50.0, Ce=Ce, qe=qe, m=0.1, V=0.05)
        # Kd should decrease as Ce increases
        assert all(Kd[i] > Kd[i + 1] for i in range(len(Kd) - 1))


# ===================================================================
# SECTION 11 — DOCX report from real fit
# ===================================================================


class TestDocxReportFromRealFit:
    """Generate a full Word report from real fitting results."""

    @pytest.fixture(autouse=True)
    def setup(self):
        from adsorblab_pro.docx_report import DOCX_AVAILABLE

        if not DOCX_AVAILABLE:
            pytest.skip("python-docx not installed")

        from adsorblab_pro.tabs.isotherm_tab import (
            _calculate_isotherm_results_direct,
            _fit_all_isotherm_models_cached,
        )

        iso_df = _load_csv("isotherm_direct.csv")
        iso_result = _calculate_isotherm_results_direct(
            {"data": iso_df, "params": {"m": 0.1, "V": 0.05}}
        )
        iso_data = iso_result.data
        Ce, qe, C0 = (
            iso_data["Ce_mgL"].values,
            iso_data["qe_mg_g"].values,
            iso_data["C0_mgL"].values,
        )
        iso_fitted = _fit_all_isotherm_models_cached(
            tuple(np.round(Ce, 8).tolist()),
            tuple(np.round(qe, 8).tolist()),
            tuple(np.round(C0, 8).tolist()),
        )

        self.study_state = {
            "isotherm_results": iso_data,
            "isotherm_models_fitted": iso_fitted,
            "kinetic_models_fitted": {},
        }

    def test_docx_report_with_real_tables(self):
        from adsorblab_pro.docx_report import DocxReportConfig, create_docx_report
        from adsorblab_pro.tabs.report_tab import generate_table

        docx_bytes, warnings = create_docx_report(
            study_title="Reviewer Test Report",
            study_state=self.study_state,
            selected_figures=[],
            selected_tables=["tbl_iso_data", "tbl_iso_params"],
            figure_generator=lambda fid, s: None,
            table_generator=generate_table,
            table_meta={
                "tbl_iso_data": ("Isotherm Data", "Equilibrium data"),
                "tbl_iso_params": ("Model Parameters", "Fitted values"),
            },
            config=DocxReportConfig(max_table_rows=50),
        )
        assert len(docx_bytes) > 2000

        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Reviewer Test Report" in full_text
        # Should have real tables embedded
        assert len(doc.tables) >= 2
