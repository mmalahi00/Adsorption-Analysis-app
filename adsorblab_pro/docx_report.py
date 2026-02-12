# docx_report.py
"""DOCX report generation for AdsorbLab Pro.

This module powers the "Word report" export in the Report tab. It is designed to be:
- Streamlit-independent (pure functions returning bytes)
- Robust to missing optional dependencies (python-docx)
- Flexible: callers provide figure/table generators, so this can be tested without Kaleido.

The report is intended for scientific publication workflows: it embeds selected figures,
key tables, and a reproducibility-oriented summary of model fits.

Export quality notes
--------------------
Plotly uses pixel-sized fonts by default. When inserted into Word at a physical width
(e.g., 6.5 inches), those fonts may appear too small. This module therefore applies
an export-only styling pass to Plotly figures (without mutating UI figures) to:
- remove in-figure headings/titles that users may add manually (captions belong in Word),
- prevent overlaps between axis titles/ticks and gradients (colorbars),
- scale typography to match the intended figure width in inches.
"""

from __future__ import annotations

from collections.abc import Callable, Mapping, Sequence
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from typing import Any

import pandas as pd

Document: Any
WD_ALIGN_PARAGRAPH: Any
Inches: Any

try:
    from docx import Document as _Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
    from docx.shared import Inches as _Inches

    Document = _Document
    WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
    Inches = _Inches
    DOCX_AVAILABLE = True
except Exception:  # pragma: no cover
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    DOCX_AVAILABLE = False


# =============================================================================
# Configuration
# =============================================================================


@dataclass(frozen=True)
class DocxReportConfig:
    """Configuration for DOCX report generation."""

    # Plotly export settings (when the figure object supports .to_image)
    img_format: str = "png"
    img_width_px: int = 1600
    img_height_px: int = 1000
    img_scale: float = 2.0

    # Document layout settings
    figure_width_in: float = 6.5

    # Typography preset for export-only Plotly styling
    # Options: "Manuscript (journal)", "Presentation (slides)", "Poster"
    text_preset: str = "Manuscript (journal)"

    # Export-only cleanups
    strip_user_headings: bool = True
    fix_axis_overlaps: bool = True

    # Table rendering limits (avoid huge DOCX)
    max_table_rows: int = 60
    max_table_cols: int = 12

    # Numeric formatting
    float_format: str = "{:.4g}"


# =============================================================================
# EXPORT FORMATTING UTILITIES - INTERNAL PACKAGE USE
# =============================================================================
# These functions are private (_name) but are re-used by:
# - adsorblab_pro.tabs.report_tab (for ZIP/PNG export formatting)
# - create_docx_report() (for DOCX report generation)
#
# They handle common Plotly export tasks:
# - Font scaling (_pt_to_px)
# - Annotation removal (_strip_user_headings_plotly)
# - Axis overlap fixing (_fix_axis_gradient_overlaps_plotly)
# - Export styling (style_figure_for_export)
#
# ⚠️ WARNING: These functions are INTERNAL and MUST NOT be imported by
# external users or external packages. They may change without notice
# between versions and have no backward compatibility guarantees.
# =============================================================================


def _pt_to_px(pt: float, *, width_px: int, target_width_in: float) -> int:
    """Convert point size to Plotly px for a given intended print width.

    Parameters
    ----------
    pt : float
        Font size in points
    width_px : int
        Export canvas width in pixels
    target_width_in : float
        Intended printed width in inches (default 6.5 for letter width)

    Returns
    -------
    int
        Font size in pixels, clamped to minimum 8px
    """
    if target_width_in <= 0:
        target_width_in = 6.5
    px = (pt / 72.0) * (float(width_px) / float(target_width_in))
    return int(max(8, round(px)))


def _strip_user_headings_plotly(fig: Any, *, fig_id: str | None = None) -> None:
    """Remove in-figure titles/headings (export-only).

    Removes:
    - Figure title text
    - Paper-level annotations (headings above plot area)

    This prevents clutter when figures are exported for manuscripts where
    captions should live in the document, not inside the image.

    Parameters
    ----------
    fig : Any
        Plotly figure object to modify in-place
    fig_id : str, optional
        Figure identifier for logging/debugging

    Notes
    -----
    Modifies the figure object in-place. This is destructive and should
    only be called on export copies, not UI figures.
    """
    # Remove heading-like annotations above plot area
    try:
        anns = list(fig.layout.annotations) if getattr(fig.layout, "annotations", None) else []
        if anns:
            kept = []
            for ann in anns:
                try:
                    xref = getattr(ann, "xref", None)
                    yref = getattr(ann, "yref", None)
                    y = getattr(ann, "y", None)
                    showarrow = getattr(ann, "showarrow", None)
                    if (
                        xref == "paper"
                        and yref == "paper"
                        and isinstance(y, (int, float))
                        and y >= 1.0
                        and showarrow is False
                    ):
                        # Skip heading-like annotations above plot area
                        continue
                    kept.append(ann)
                except Exception:
                    kept.append(ann)
            fig.update_layout(annotations=kept)
    except Exception:
        pass

    # Clear layout title text
    try:
        title_text = ""
        try:
            title_text = str(fig.layout.title.text or "")
        except Exception:
            title_text = ""
        if title_text.strip():
            fig.update_layout(title_text="")
    except Exception:
        pass


def _fix_axis_gradient_overlaps_plotly(
    fig: Any,
    *,
    width_px: int,
    height_px: int,
    tick_px: int,
    axis_title_px: int,
    font_family: str = "Arial, sans-serif",
) -> None:
    """Prevent overlaps between axes, ticks, and gradients (export-only).

    Parameters
    ----------
    fig : Any
        Plotly figure object to modify in-place
    width_px : int
        Export canvas width in pixels
    height_px : int
        Export canvas height in pixels
    tick_px : int
        Tick label font size in pixels
    axis_title_px : int
        Axis title font size in pixels
    font_family : str
        Font family for colorbar labels

    Notes
    -----
    Adjusts margins and colorbar positioning to prevent overlaps when
    figures are embedded in documents at specified print widths.
    """
    # Always increase standoff + automargins
    try:
        standoff = max(24, int(1.4 * axis_title_px))
        fig.update_xaxes(automargin=True, title_standoff=standoff)
        fig.update_yaxes(automargin=True, title_standoff=standoff)
    except Exception:
        pass

    # Base safe margins (even without colorbars)
    try:
        has_title = bool(str(getattr(fig.layout.title, "text", "") or "").strip())
    except Exception:
        has_title = False

    try:
        cur = fig.layout.margin.to_plotly_json() if getattr(fig.layout, "margin", None) else {}
        margin_l = int(cur.get("l", 70))
        r = int(cur.get("r", 40))
        t = int(cur.get("t", 40))
        b = int(cur.get("b", 70))
    except Exception:
        margin_l, r, t, b = 70, 40, 40, 70

    min_l = max(90, int(0.09 * width_px), int(2.8 * axis_title_px), int(2.4 * tick_px))
    min_b = max(90, int(0.10 * height_px), int(2.4 * axis_title_px), int(2.0 * tick_px))
    min_r = max(60, int(0.05 * width_px), int(1.8 * tick_px))
    min_t = (
        max(55, int(0.06 * height_px), int(1.2 * axis_title_px))
        if has_title
        else max(25, int(0.03 * height_px))
    )

    margin_l = max(margin_l, min_l)
    b = max(b, min_b)
    r = max(r, min_r)
    t = max(t, min_t)

    try:
        fig.update_layout(margin={"l": margin_l, "r": r, "t": t, "b": b, "pad": 10})
    except Exception:
        pass

    # Detect colorbars and, if present, push outward + add margin
    try:
        fig_json = fig.to_plotly_json()
        traces = fig_json.get("data", []) or []
        layout = fig_json.get("layout", {}) or {}

        has_colorbar = False
        colorbar_is_horizontal = False

        for tr in traces:
            if not isinstance(tr, dict):
                continue

            cb = tr.get("colorbar")
            if isinstance(cb, dict):
                has_colorbar = True
                if str(cb.get("orientation", "")).lower().startswith("h"):
                    colorbar_is_horizontal = True
                break

            marker = tr.get("marker")
            if isinstance(marker, dict) and isinstance(marker.get("colorbar"), dict):
                has_colorbar = True
                cb2 = marker.get("colorbar") or {}
                if str(cb2.get("orientation", "")).lower().startswith("h"):
                    colorbar_is_horizontal = True
                break

            line = tr.get("line")
            if isinstance(line, dict) and isinstance(line.get("colorbar"), dict):
                has_colorbar = True
                cb3 = line.get("colorbar") or {}
                if str(cb3.get("orientation", "")).lower().startswith("h"):
                    colorbar_is_horizontal = True
                break

        if not has_colorbar:
            for k, v in layout.items():
                if (
                    str(k).startswith("coloraxis")
                    and isinstance(v, dict)
                    and isinstance(v.get("colorbar"), dict)
                ):
                    has_colorbar = True
                    cb4 = v.get("colorbar") or {}
                    if str(cb4.get("orientation", "")).lower().startswith("h"):
                        colorbar_is_horizontal = True
                    break

        if not has_colorbar:
            return

        # Extra margin for colorbar
        try:
            cur = fig.layout.margin.to_plotly_json() if getattr(fig.layout, "margin", None) else {}
            margin_l = int(cur.get("l", margin_l))
            r = int(cur.get("r", r))
            t = int(cur.get("t", t))
            b = int(cur.get("b", b))

            if colorbar_is_horizontal:
                b = max(b, int(max(140, 0.12 * height_px)))
            else:
                r = max(r, int(max(180, 0.14 * width_px)))

            fig.update_layout(margin={"l": margin_l, "r": r, "t": t, "b": b, "pad": 10})
        except Exception:
            pass

        # Push layout-level coloraxis colorbars outward
        try:
            coloraxis_keys = [k for k in layout.keys() if str(k).startswith("coloraxis")]
            for k in coloraxis_keys:
                v = layout.get(k, {})
                if isinstance(v, dict) and isinstance(v.get("colorbar"), dict):
                    fig.update_layout(
                        {
                            k: {
                                "colorbar": {
                                    "x": 1.05,
                                    "xanchor": "left",
                                    "xpad": 14,
                                    "ypad": 10,
                                    "tickfont": {"size": tick_px, "family": font_family},
                                }
                            }
                        }
                    )
        except Exception:
            pass

        # Push trace-level colorbars outward
        try:
            for tr in fig.data:
                if getattr(tr, "colorbar", None) is not None:
                    tr.colorbar.update(
                        x=1.05,
                        xanchor="left",
                        xpad=14,
                        ypad=10,
                        tickfont={"size": tick_px, "family": font_family},
                    )
                if (
                    getattr(tr, "marker", None) is not None
                    and getattr(tr.marker, "colorbar", None) is not None
                ):
                    tr.marker.colorbar.update(
                        x=1.05,
                        xanchor="left",
                        xpad=14,
                        ypad=10,
                        tickfont={"size": tick_px, "family": font_family},
                    )
        except Exception:
            pass

    except Exception:
        pass


def style_figure_for_export(
    fig: Any,
    *,
    width_px: int = 1600,
    height_px: int = 1000,
    tick_px: int = 12,
) -> None:
    """
    Style a Plotly figure for export in ZIP and other formats.

    This is a convenience wrapper around export utilities that applies
    common export fixes in a single call:
    - Removes user-added headings and titles
    - Fixes axis/gradient label overlaps
    - Scales fonts appropriately

    Modifies figure in-place.

    Parameters
    ----------
    fig : Any
        Plotly figure object
    width_px : int
        Export width in pixels (default 1600)
    height_px : int
        Export height in pixels (default 1000)
    tick_px : int
        Tick label size in pixels (default 12)

    Notes
    -----
    This function is used by adsorblab_pro.tabs.report_tab for
    ZIP export formatting. It provides a simpler interface than
    _style_plotly_for_docx_export() which requires DocxReportConfig.
    """
    _fix_axis_gradient_overlaps_plotly(
        fig,
        width_px=width_px,
        height_px=height_px,
        tick_px=tick_px,
        axis_title_px=12,
        font_family="Arial, sans-serif",
    )
    _strip_user_headings_plotly(fig, fig_id=None)


def _style_plotly_for_docx_export(fig_obj: Any, cfg: DocxReportConfig) -> Any:
    """Return a copy of a Plotly figure with DOCX-friendly export styling."""
    # Local imports: keep module import-light for non-Plotly use
    try:
        import plotly.graph_objects as go
    except ImportError:
        return fig_obj  # pragma: no cover

    if fig_obj is None:
        return fig_obj

    # Copy: don't mutate caller's figure
    fig = go.Figure(fig_obj)

    # STEP 1: Remove user headings (if config allows)
    if cfg.strip_user_headings:
        _strip_user_headings_plotly(fig)

    # STEP 2: Force export dimensions
    try:
        fig.update_layout(width=int(cfg.img_width_px), height=int(cfg.img_height_px))
    except Exception:
        pass

    # STEP 3: Typography presets
    preset_norm = (cfg.text_preset or "Manuscript (journal)").strip().lower()
    if preset_norm.startswith("present"):
        base_pt, tick_pt, axis_title_pt, title_pt, legend_pt = 16.0, 15.0, 17.0, 20.0, 14.0
    elif preset_norm.startswith("poster"):
        base_pt, tick_pt, axis_title_pt, title_pt, legend_pt = 18.0, 18.0, 20.0, 24.0, 16.0
    else:  # Manuscript (default)
        base_pt, tick_pt, axis_title_pt, title_pt, legend_pt = 11.0, 10.5, 12.0, 14.0, 10.0

    base_px = _pt_to_px(base_pt, width_px=cfg.img_width_px, target_width_in=cfg.figure_width_in)
    tick_px = _pt_to_px(tick_pt, width_px=cfg.img_width_px, target_width_in=cfg.figure_width_in)
    axis_title_px = _pt_to_px(
        axis_title_pt, width_px=cfg.img_width_px, target_width_in=cfg.figure_width_in
    )
    title_px = _pt_to_px(title_pt, width_px=cfg.img_width_px, target_width_in=cfg.figure_width_in)
    legend_px = _pt_to_px(legend_pt, width_px=cfg.img_width_px, target_width_in=cfg.figure_width_in)

    # Base font
    try:
        fig.update_layout(font={"family": "Arial, sans-serif", "size": base_px})
    except Exception:
        pass

    # 2D axes typography
    try:
        fig.update_xaxes(
            tickfont={"size": tick_px, "family": "Arial, sans-serif"},
            title_font={"size": axis_title_px, "family": "Arial, sans-serif"},
        )
        fig.update_yaxes(
            tickfont={"size": tick_px, "family": "Arial, sans-serif"},
            title_font={"size": axis_title_px, "family": "Arial, sans-serif"},
        )
    except Exception:
        pass

    # Legend
    try:
        if getattr(fig.layout, "legend", None) is not None:
            fig.update_layout(legend={"font": {"size": legend_px, "family": "Arial, sans-serif"}})
    except Exception:
        pass

    # Legend symbols: constant size
    try:
        leg = fig.layout.legend.to_plotly_json() if getattr(fig.layout, "legend", None) else {}
        leg.update({"itemsizing": "constant"})
        fig.update_layout(legend=leg)
    except Exception:
        pass

    # Title font
    try:
        if getattr(fig.layout, "title", None) is not None:
            current_text = ""
            try:
                current_text = str(fig.layout.title.text or "")
            except Exception:
                current_text = ""
            if current_text.strip() != "":
                fig.update_layout(title={"font": {"size": title_px, "family": "Arial, sans-serif"}})
    except Exception:
        pass

    # Annotations: make typography consistent
    try:
        if getattr(fig.layout, "annotations", None):
            for ann in fig.layout.annotations:
                try:
                    ann_font = ann.font.to_plotly_json() if getattr(ann, "font", None) else {}
                    ann_font.update({"family": "Arial, sans-serif", "size": base_px})
                    ann.font = ann_font
                except Exception:
                    pass
    except Exception:
        pass

    # Prevent axis-title vs gradient/colorbar overlaps
    if cfg.fix_axis_overlaps:
        _fix_axis_gradient_overlaps_plotly(
            fig,
            width_px=cfg.img_width_px,
            height_px=cfg.img_height_px,
            tick_px=tick_px,
            axis_title_px=axis_title_px,
            font_family="Arial, sans-serif",
        )

    return fig


# =============================================================================
# Helpers for document construction
# =============================================================================


def _ensure_caption_style(doc: Any) -> str | None:
    """Ensure 'Caption' style exists in document; return its name or None."""
    try:
        return "Caption"
    except Exception:
        pass
    return None


def _add_dataframe_table(
    doc: Any,
    df: pd.DataFrame,
    caption: str,
    config: DocxReportConfig,
) -> None:
    """Insert a table from a DataFrame."""
    # Truncate rows if needed
    if len(df) > config.max_table_rows:
        df = df.iloc[: config.max_table_rows]

    # Truncate columns if needed
    if len(df.columns) > config.max_table_cols:
        df = df.iloc[:, : config.max_table_cols]

    # Create table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            if pd.isna(val):
                row_cells[i].text = ""
            else:
                try:
                    if isinstance(val, float):
                        row_cells[i].text = config.float_format.format(val)
                    else:
                        row_cells[i].text = str(val)
                except Exception:
                    row_cells[i].text = str(val)

    # Add caption
    cap_style = _ensure_caption_style(doc)
    if cap_style:
        cap_par = doc.add_paragraph(caption, style=cap_style)
    else:
        cap_par = doc.add_paragraph(caption)
        try:
            cap_par.runs[0].italic = True
        except Exception:
            pass
    try:
        cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    doc.add_paragraph(" ")  # spacer


def _coerce_image_bytes(fig_obj: Any, config: DocxReportConfig) -> bytes:
    """Convert Plotly figure to image bytes."""
    # Apply export styling (copy, don't mutate)
    styled_fig = _style_plotly_for_docx_export(fig_obj, config)

    # Try Plotly's native export
    try:
        return styled_fig.to_image(
            format=config.img_format,
            width=config.img_width_px,
            height=config.img_height_px,
            scale=config.img_scale,
        )
    except Exception:
        pass

    # Fallback: use kaleido if available
    try:
        import kaleido  # noqa: F401

        return styled_fig.to_image(
            format=config.img_format,
            width=config.img_width_px,
            height=config.img_height_px,
            scale=config.img_scale,
        )
    except Exception:
        raise RuntimeError("Cannot export figure to image. Install kaleido: pip install kaleido")


def _add_figure(
    doc: Any,
    fig_bytes: bytes,
    caption: str,
    config: DocxReportConfig,
) -> None:
    """Insert an image and caption."""
    bio = BytesIO(fig_bytes)
    bio.seek(0)
    doc.add_picture(bio, width=Inches(config.figure_width_in))
    # Center the paragraph containing the picture
    try:
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    cap_style = _ensure_caption_style(doc)
    if cap_style:
        cap_par = doc.add_paragraph(caption, style=cap_style)
    else:
        cap_par = doc.add_paragraph(caption)
        try:
            cap_par.runs[0].italic = True
        except Exception:
            pass
    try:
        cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    doc.add_paragraph(" ")  # spacer


def _best_model_line(models: Any, label: str) -> str | None:
    """Return a short 'best model' summary line if possible."""
    if not isinstance(models, dict):
        return None
    candidates = []
    for name, payload in models.items():
        if isinstance(payload, dict) and payload.get("converged"):
            r2 = payload.get("r_squared")
            try:
                r2v = float(r2) if r2 is not None else float("-inf")
            except Exception:
                r2v = float("-inf")
            candidates.append((r2v, name))
    if not candidates:
        return None
    best_r2, best_name = max(candidates, key=lambda x: x[0])
    if best_r2 == float("-inf"):
        return None
    return f"{label}: {best_name} (R²={best_r2:.4f})"


# =============================================================================
# Public API
# =============================================================================


def create_docx_report(
    *,
    study_title: str,
    study_state: Mapping[str, Any],
    selected_figures: Sequence[str],
    selected_tables: Sequence[str],
    figure_generator: Callable[[str, Mapping[str, Any]], Any],
    table_generator: Callable[[str, Mapping[str, Any]], pd.DataFrame | None],
    figure_meta: Mapping[str, tuple[str, str]] | None = None,
    table_meta: Mapping[str, tuple[str, str]] | None = None,
    config: DocxReportConfig | None = None,
) -> tuple[bytes, list[str]]:
    """Create a publication-oriented Word report and return it as bytes.

    Returns:
        (docx_bytes, warnings)
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")

    cfg = config or DocxReportConfig()
    warnings: list[str] = []

    doc = Document()

    # Title
    doc.add_heading(study_title, level=0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Produced by AdsorbLab Pro (Report Export)")
    doc.add_paragraph(" ")

    # Study overview / reproducibility
    doc.add_heading("Study overview", level=1)
    overview = doc.add_paragraph()
    overview.add_run("Included analyses: ").bold = True

    present = []
    if study_state.get("isotherm_models_fitted") is not None:
        present.append("Isotherms")
    if study_state.get("kinetic_models_fitted") is not None:
        present.append("Kinetics")
    if study_state.get("thermo_params") is not None:
        present.append("Thermodynamics")
    if study_state.get("ph_effect_results") is not None:
        present.append("pH effect")
    if study_state.get("temp_effect_results") is not None:
        present.append("Temperature effect")
    if study_state.get("dosage_results") is not None:
        present.append("Dosage effect")
    overview.add_run(", ".join(present) if present else "(no analyses detected)")

    # Best-model summaries (if available)
    summary_lines = []
    bm = _best_model_line(study_state.get("isotherm_models_fitted"), "Best isotherm model")
    if bm:
        summary_lines.append(bm)
    bm = _best_model_line(study_state.get("kinetic_models_fitted"), "Best kinetic model")
    if bm:
        summary_lines.append(bm)

    thermo = study_state.get("thermo_params")
    if isinstance(thermo, dict) and thermo.get("success"):
        try:
            dH = thermo.get("delta_H")
            dS = thermo.get("delta_S")
            dG = thermo.get("delta_G")
            summary_lines.append(
                "Thermodynamics: ΔH={:.4g}, ΔS={:.4g}, ΔG={}".format(
                    float(dH) if dH is not None else float("nan"),
                    float(dS) if dS is not None else float("nan"),
                    str(dG),
                )
            )
        except Exception:
            summary_lines.append("Thermodynamics: (computed; see tables section)")

    if summary_lines:
        doc.add_paragraph(" ")
        doc.add_heading("Key results", level=2)
        for line in summary_lines:
            doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph(" ")

    # Methods (brief, publication-friendly)
    doc.add_heading("Methods summary", level=1)
    methods = [
        "Nonlinear regression is used to fit adsorption models when possible.",
        "Model quality may include metrics such as R², adjusted R², AIC/AICc, BIC, RMSE, and residual diagnostics (depending on the analysis).",
        "Figures are exported as high-resolution static images for manuscript preparation.",
        "Tables are exported in a format suitable for copy/paste into manuscripts or supplementary information.",
    ]
    for line in methods:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph(" ")

    # Tables section
    doc.add_heading("Tables", level=1)
    if not selected_tables:
        doc.add_paragraph(
            "No tables were selected. Tip: select parameter and comparison tables to include them in this report."
        )
    else:
        for idx, tbl_id in enumerate(selected_tables, start=1):
            df = None
            try:
                df = table_generator(tbl_id, study_state)
            except Exception as e:
                warnings.append(f"Table '{tbl_id}' could not be generated: {e}")
            if df is None or df.empty:
                warnings.append(f"Table '{tbl_id}' is empty or unavailable.")
                continue

            title = tbl_id
            desc = ""
            if table_meta and tbl_id in table_meta:
                title, desc = table_meta[tbl_id]
            caption = f"Table {idx}. {title}" + (f" — {desc}" if desc else "")
            _add_dataframe_table(doc, df, caption, cfg)

    doc.add_paragraph(" ")

    # Figures section
    doc.add_heading("Figures", level=1)
    if not selected_figures:
        doc.add_paragraph(
            "No figures were selected. Tip: select key plots (fits, residuals, diagnostics) to include them in this report."
        )
    else:
        for idx, fig_id in enumerate(selected_figures, start=1):
            fig_obj = None
            try:
                fig_obj = figure_generator(fig_id, study_state)
            except Exception as e:
                warnings.append(f"Figure '{fig_id}' could not be generated: {e}")
                continue

            if fig_obj is None:
                warnings.append(f"Figure '{fig_id}' is unavailable.")
                continue

            try:
                img_bytes = _coerce_image_bytes(fig_obj, cfg)
            except Exception as e:
                warnings.append(f"Figure '{fig_id}' export failed: {e}")
                continue

            title = fig_id
            desc = ""
            if figure_meta and fig_id in figure_meta:
                title, desc = figure_meta[fig_id]
            caption = f"Figure {idx}. {title}" + (f" — {desc}" if desc else "")
            _add_figure(doc, img_bytes, caption, cfg)

    # Notes / warnings
    if warnings:
        doc.add_heading("Notes", level=1)
        doc.add_paragraph("Some items could not be included (e.g., missing data or export issues):")
        for w in warnings[:30]:
            doc.add_paragraph(w, style="List Bullet")
        if len(warnings) > 30:
            doc.add_paragraph(f"… and {len(warnings) - 30} more", style="List Bullet")

    # Serialize
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue(), warnings
